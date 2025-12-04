import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import cv2
import numpy as np
from typing import List, Tuple, Optional
from modules.config import config
from modules.utils import get_logger

logger = get_logger(__name__)


class OCRProcessor:
    """Handles OCR operations for different file types

    This class provides functionality to extract text from various file formats
    including PDF, images (JPG, PNG), Word documents and Excel spreadsheets
    using both direct extraction methods and OCR technology.
    """

    def __init__(self):
        """Initialize the OCR processor with configuration settings."""
        # Set Tesseract path
        if config.TESSERACT_PATH:
            pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_PATH
        # Set Tesseract data directory
        if config.TESSDATA_PREFIX:
            pytesseract.tessdata_dir = config.TESSDATA_PREFIX

    def preprocess_image(self, image: Image.Image, enhance_contrast: bool = True, dpi: int = 300) -> Image.Image:
        """Preprocess image to improve OCR accuracy with advanced techniques

        Args:
            image (PIL.Image.Image): Input image to preprocess
            enhance_contrast (bool): Whether to enhance image contrast
            dpi (int): Resolution for image processing

        Returns:
            PIL.Image.Image: Preprocessed image ready for OCR
        """
        try:
            # Convert to OpenCV format for advanced preprocessing
            img_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

            # Convert to grayscale if not already
            if len(img_cv.shape) == 3:
                gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
            else:
                gray = img_cv

            # Increase image resolution if needed
            height, width = gray.shape
            if height < 300 or width < 300:
                # Scale up the image to improve OCR accuracy
                scale_factor = max(300 / height, 300 / width)
                new_width = int(width * scale_factor)
                new_height = int(height * scale_factor)
                gray = cv2.resize(gray, (new_width, new_height), interpolation=cv2.INTER_CUBIC)

            # Apply Gaussian blur to reduce noise
            gray = cv2.GaussianBlur(gray, (3, 3), 0)

            # Enhance contrast if requested
            if enhance_contrast:
                # Apply adaptive histogram equalization (CLAHE)
                clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
                gray = clahe.apply(gray)

            # Try multiple thresholding methods and choose the best one
            # Otsu's method
            _, thresh_otsu = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

            # Adaptive threshold
            adaptive_thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)

            # Choose the thresholding method that gives better results
            # For documents, Otsu's method usually works well
            processed = thresh_otsu

            # Additional morphological operations to clean up the image
            # Remove small noise
            kernel = np.ones((1, 1), np.uint8)
            processed = cv2.morphologyEx(processed, cv2.MORPH_CLOSE, kernel)
            processed = cv2.morphologyEx(processed, cv2.MORPH_OPEN, kernel)

            # Apply morphological operations to connect broken text
            kernel = np.ones((2, 2), np.uint8)
            processed = cv2.dilate(processed, kernel, iterations=1)

            # Deskew the image if needed
            processed = self._deskew_image(processed)

            # Convert back to PIL image
            processed_image = Image.fromarray(processed)

            return processed_image
        except Exception as e:
            logger.error(f"Error preprocessing image: {e}")
            # Return original image as fallback
            return image

    def _deskew_image(self, image):
        """
        Deskew the image to correct rotation for better OCR results
        """
        coords = np.column_stack(np.where(image > 0))

        if coords.size == 0:
            # If no foreground pixels found, return original image
            return image

        angle = cv2.minAreaRect(coords)[-1]

        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle

        (h, w) = image.shape[:2]
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, angle, 1.0)
        rotated = cv2.warpAffine(image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

        return rotated

    def extract_text_from_pdf(self, pdf_stream) -> str:
        """
        Extract text from PDF using direct extraction or OCR as fallback
        """
        try:
            doc = fitz.open(stream=pdf_stream, filetype="pdf")
            all_text = []

            for page_num in range(len(doc)):
                try:
                    page = doc.load_page(page_num)
                    text_from_page = page.get_text("text")

                    if text_from_page.strip():
                        # If there's text directly in the PDF, use it
                        all_text.append(text_from_page)
                    else:
                        # If no text, use OCR on the page image
                        pix = page.get_pixmap(dpi=config.DEFAULT_DPI)
                        img_bytes = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_bytes))

                        # Preprocess the image for better OCR
                        processed_img = self.preprocess_image(img,
                                                            enhance_contrast=config.CONTRAST_ENHANCEMENT,
                                                            dpi=config.DEFAULT_DPI)

                        ocr_text = pytesseract.image_to_string(processed_img, lang='spa')
                        all_text.append(ocr_text)
                except Exception as e:
                    logger.error(f"Error processing page {page_num} of PDF: {e}")
                    continue

            doc.close()
            return "\n\n--- Página Siguiente ---\n\n".join(all_text)
        except Exception as e:
            logger.error(f"Error opening or processing PDF: {e}")
            raise

    def extract_text_from_image(self, image: Image.Image, preprocess: bool = True) -> str:
        """
        Extract text from image using OCR
        """
        try:
            if preprocess:
                processed_img = self.preprocess_image(image,
                                                    enhance_contrast=config.CONTRAST_ENHANCEMENT,
                                                    dpi=config.DEFAULT_DPI)
            else:
                processed_img = image

            # Perform OCR
            text = pytesseract.image_to_string(processed_img, lang='spa')
            return text
        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            raise

    def process_file(self, file_stream, file_type: str, preprocess: bool = True) -> str:
        """
        Main method to process files based on type
        """
        try:
            # Validate inputs
            if not file_stream:
                raise ValueError("File stream cannot be empty")
            if not file_type:
                raise ValueError("File type cannot be empty")
            if not isinstance(file_type, str):
                raise ValueError("File type must be a string")

            # Validate supported file type
            file_type_lower = file_type.lower()
            supported_types = ['pdf', 'jpg', 'jpeg', 'png', 'docx', 'xlsx', 'xls']
            if file_type_lower not in supported_types:
                error_msg = f"Unsupported file type: {file_type}. Supported types: {supported_types}"
                logger.error(error_msg)
                raise ValueError(error_msg)

            # Calculate file hash for caching (need to handle both bytes and file-like objects)
            from modules.utils import calculate_file_hash, load_cached_result, save_processed_result

            # Check if file_stream is bytes, convert to BytesIO if needed for hash calculation
            if isinstance(file_stream, bytes):
                temp_file = io.BytesIO(file_stream)
                file_hash = calculate_file_hash(temp_file)
            else:
                file_hash = calculate_file_hash(file_stream)

            # Check if result is already cached
            cached_result = load_cached_result(file_hash)
            if cached_result:
                logger.info(f"Using cached result for file with hash {file_hash}")
                return cached_result["extracted_text"]

            # Process the file normally
            if file_type_lower == 'pdf':
                # Handle both bytes and file-like objects for PDF
                if isinstance(file_stream, bytes):
                    result = self.extract_text_from_pdf(io.BytesIO(file_stream))
                else:
                    result = self.extract_text_from_pdf(file_stream)
            elif file_type_lower in ['jpg', 'jpeg', 'png']:
                if isinstance(file_stream, bytes):
                    image = Image.open(io.BytesIO(file_stream))
                else:
                    image = Image.open(file_stream)
                result = self.extract_text_from_image(image, preprocess)
            elif file_type_lower == 'docx':
                # Handle both bytes and file-like objects for DOCX
                if isinstance(file_stream, bytes):
                    result = self.extract_text_from_docx(io.BytesIO(file_stream))
                else:
                    result = self.extract_text_from_docx(file_stream)
            elif file_type_lower in ['xlsx', 'xls']:
                # Handle both bytes and file-like objects for Excel
                if isinstance(file_stream, bytes):
                    result = self.extract_text_from_excel(io.BytesIO(file_stream))
                else:
                    result = self.extract_text_from_excel(file_stream)
            else:
                # This case should not occur due to validation above, but added for safety
                error_msg = f"Unsupported file type: {file_type}"
                logger.error(error_msg)
                raise ValueError(error_msg)

            # Cache the result
            save_processed_result(file_hash, result, {})

            return result
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Error processing file of type {file_type}: {e}")
            raise

    def extract_text_from_docx(self, file_stream) -> str:
        """
        Extract text from DOCX file
        """
        try:
            from docx import Document
            import io

            # Create a temporary file-like object from the stream
            file_stream.seek(0)
            temp_docx = io.BytesIO(file_stream.read())
            doc = Document(temp_docx)

            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)

            return "\n\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise

    def extract_text_from_excel(self, file_stream) -> str:
        """
        Extract text from Excel (XLS, XLSX) file
        """
        try:
            import pandas as pd
            import io

            # Create a temporary file-like object from the stream
            file_stream.seek(0)
            temp_xlsx = io.BytesIO(file_stream.read())

            # Read all sheets
            excel_file = pd.ExcelFile(temp_xlsx)
            text_parts = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text_parts.append(f"--- Hoja: {sheet_name} ---")
                text_parts.append(df.to_string())

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from Excel: {e}")
            raise